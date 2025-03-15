import scapy.all as scapy
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from datetime import datetime
import os
from collections import Counter


def analyze_packet(packet):
    src_ip = None
    dst_ip = None
    protocol = None

    if packet.haslayer(scapy.IP):
        src_ip = packet[scapy.IP].src
        dst_ip = packet[scapy.IP].dst
        protocol = "IP"
    elif packet.haslayer(scapy.IPv6):
        src_ip = packet[scapy.IPv6].src
        dst_ip = packet[scapy.IPv6].dst
        protocol = "IPv6"

    if packet.haslayer(scapy.TCP):
        protocol += " TCP"
    elif packet.haslayer(scapy.UDP):
        protocol += " UDP"
    elif packet.haslayer(scapy.ICMP):
        protocol += " ICMP"

    return {
        'time': datetime.now(),
        'src_ip': src_ip,
        'dst_ip': dst_ip,
        'protocol': protocol,
        'length': len(packet)
    }


def is_suspicious(packet_info):
    if packet_info['protocol'] == "IP TCP" and packet_info['length'] >= 1500:
        return True
    if packet_info['protocol'] == "IP UDP" and packet_info['length'] >= 1000:
        return True
    return False


def is_alarming(packet_info):
    if packet_info['protocol'] == "IP TCP" and packet_info['length'] >= 1000:
        return True
    if packet_info['protocol'] == "IP UDP" and packet_info['length'] >= 800:
        return True
    return False


def create_report(packet_data):
    document = Document()

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)

    heading = document.add_paragraph()
    heading_run = heading.add_run("Отчет об анализе сетевого трафика")
    heading_run.font.size = Pt(16)
    heading_run.bold = True
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph(f"Дата и время создания отчета: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}")
    document.add_paragraph(f"")

    packet_counter = Counter(packet['src_ip'] for packet in packet_data if packet['src_ip'] is not None)

    top_hosts = packet_counter.most_common(10)

    document.add_paragraph("Топ 10 хостов с наибольшим количеством пакетов:")
    for host, count in top_hosts:
        document.add_paragraph(f"Хост: {host}, Количество пакетов: {count}")

    document.add_paragraph(f"")  # Пустая строка

    for entry in packet_data:
        paragraph = document.add_paragraph()
        time_str = entry['time'].strftime('%H:%M:%S')
        text_run = paragraph.add_run(f"Время: {time_str} | ")
        text_run.bold = True
        paragraph.add_run(
            f"Src IP: {entry['src_ip']} | Dst IP: {entry['dst_ip']} | Протокол: {entry['protocol']} | Размер: {entry['length']} ")

        if is_suspicious(entry):
            text_run.font.color.rgb = RGBColor(255, 0, 0)  # красный
        elif is_alarming(entry):
            text_run.font.highlight_color = 15  # желтый
        else:
            text_run.font.color.rgb = RGBColor(0, 128, 0)  # зеленый

    file_name = f"network_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    document.save(file_name)

    return os.path.dirname(os.path.abspath(file_name))
